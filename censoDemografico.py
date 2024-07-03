import collections
import glob
import io
import os
import subprocess
import shutil
import re
import zipfile

import docx
import pandas as pd
from pandas.api.types import CategoricalDtype
import pyreadstat
import py7zr
import requests

from .common import handleDatabase, mean_weight, std_weight, median_weight,\
                    print_info, print_error, parse_sas
from .definitions import FILETYPES_PATH, RAW_FILES_PATH, UF_SIGLA_NOME, SUPPORTED_FTs

PATH = 'censo-demografico'

URL = {
    1960: 'https://drive.google.com/uc?export=download&id=1ehlPo10QweI9xCj_3L6QnYRNfEnGP1nv',
    1970: 'https://drive.usercontent.google.com/download?id=1lcvKVIuBYczyx31CB7tRhyHxHcdjqb4A&export=download&authuser=0&confirm=t&uuid=a3100044-2435-4d25-85ce-9fb5b781c30b&at=APZUnTV0fb2psgi8DTGdzjna9YcP%3A1714616660467',
    1980: 'https://drive.usercontent.google.com/download?id=1gOCtxn9rbGTzvBzHDIf7iseljTnBgBH2&export=download&authuser=0&confirm=t&uuid=a919e85e-a488-445a-936c-3e9458d53e96&at=APZUnTWPnlNZOt3fnD61ovUYjP1j%3A1714687784460',
    1991: 'https://drive.usercontent.google.com/download?id=1T3yAwWwkqDZ4K-macO0bDg42YKh2juKF&export=download&authuser=0&confirm=t&uuid=3c130385-83bd-4130-a70e-cb79250c6194&at=APZUnTXmK7duBJ3DOcDagyXGaSQ2%3A1714688091118',
    2000: 'https://ftp.ibge.gov.br/Censos/Censo_Demografico_2000/Microdados',
    2010: 'https://ftp.ibge.gov.br/Censos/Censo_Demografico_2010/Resultados_Gerais_da_Amostra/Microdados'}

TYPES = ('PESS', 'DOMI')

EXPR_FILTER = {
    'ALL': ('[file_url["href"] for file_url in soup.find_all("a")'
                               ' if "zip" in file_url["href"]]'),
    '_' :('[file_url["href"] for file_url in soup.find_all("a")'
                               ' if "zip" in file_url["href"]' 
                               ' and f"{self.uf}" in file_url["href"]]')
}

DOCUMENTACAO = {
    'PESS': {1960: '', #TODO
             1970: 'Censo 1970/Documentação/Amostra 1970 vol03.doc',
             1980: '', #TODO
             1991: '', #TODO
             2000: 'LE PESSOAS.sas',
             2010: 'Layout_microdados_Amostra.xls'},
    'DOMI': {1970: 'Censo 1970/Documentação/Amostra 1970 vol03.doc',
             1980: '', #TODO
             1991: '', #TODO
             2000: 'LE DOMIC.sas',
             2010: 'Layout_microdados_Amostra.xls'}
}

WEIGHTS = {
    1960: None,
    1970: 'V054',
    1980: 'V604',
    1991: 'V7301',
    2000: 'P001',
    2010: 'V0010',
}

RAW_FILENAME = {
    1960: 'Censo Demográfico de 1960.7z',
    1970: 'Censo Demográfico de 1970.7z',
    1980: 'Censo Demográfico de 1980.7z',
    1991: 'Censo Demográfico de 1991.7z',
}

VARS = {
    'sexo': {
             1960: None,
             1970: 'V054',
             1980: '', #TODO
             1991: '', #TODO
             2000: 'P001',
             2010: 'V0010',
    },
}

class handleCensoDemografico(handleDatabase):
    def __init__(self, year, uf, type_db, medium=requests):
        if year not in (1960, 1970, 1980, 1991, 2000, 2010):
            print_error(f'Ano {year} não implementado.')
            raise ValueError 
        if uf not in UF_SIGLA_NOME and uf.upper() != 'ALL':
            print_error(f'UF {uf} não implementada. As opções válidas são'
                        f'{UF_SIGLA_NOME.keys()} e "all"')
            raise ValueError
        if year <= 1991 and uf.upper() != 'ALL':
            print_error('Só estão disponíveis arquivos do Censo Demográfico '
                        'de 1960 e 1991 de todo o Brasil. Por favor, utilize '
                        'a opção "ALL"')
            raise ValueError

        if type_db not in TYPES:
            print_error(f'Tipo {type_db} não existente. As opções válidas são'
                        f'{TYPES.keys()}')
            raise ValueError

        self.type_db = type_db
        self.uf = uf.upper()
        super().__init__(year, medium)
        self.filename = f'{self.year}-{self.type_db}-{self.uf}-censo-demografico'
        self.name = 'Censo Demográfico'
        self.doc_filename = DOCUMENTACAO[self.type_db][self.year]
        self.weight_var = WEIGHTS[self.year] 
        self.path = os.path.join(self.root, PATH)
        if not os.path.isdir(self.path):
            os.mkdir(self.path)
        self.path = os.path.join(self.path, str(year))
        self.path_dict = os.path.join(self.path, f'{self.type_db}-dicionario') 
        if not os.path.isdir(self.path):
            os.mkdir(self.path)
        self.raw_files_path = os.path.join(self.path, RAW_FILES_PATH)
        if not os.path.isdir(self.raw_files_path):
            os.mkdir(self.raw_files_path)
        if self.year < 2000:
            self.raw_filename = RAW_FILENAME[self.year]
            self.raw_filepath = os.path.join(self.raw_files_path, self.raw_filename)
        self.url = URL[year]
        self.expr_filter = EXPR_FILTER[self.uf if self.uf == 'ALL' else '_']
        self.is_zipped = True

    def get_url(self):
        if self.year < 2000:
            self.file_urls = [self.url]
        else:
            file_urls = super().get_url(criterion, unique=False)
            self.file_urls = [os.path.join(self.url, file_url)
                              for file_url in file_urls]
        return self.file_urls

    def get_save_raw_database(self):
        self.get_url()
        self.filepaths = []
        for file_url in self.file_urls:
            filepath = super().get_save_raw_database(file_url)
            self.filepaths.append(filepath)

    def make_database_dict(self):
        docpath = glob.glob(f'{self.raw_files_path}/*[Dd]oc*.zip')[0]

        with zipfile.ZipFile(docpath, metadata_encoding='cp850') as zf:
            for file_path in zf.namelist():
                filename = os.path.split(file_path)[-1]
                if filename == self.doc_filename:
                    fp = file_path
                    break
            try:
                print_info(f'Extraíndo informações do arquivo {fp}')
            except UnboundLocalError:
                print_error(f'Não foi possível encontrar o arquivo {self.doc_filename}')
                raise UnboundLocalError
            with zf.open(fp) as f:
                match self.year:
                    case 2000:
                        parse_sas(self, f, encoding='latin-1')
                    case 2010:
                        self.df_dict = pd.read_excel(f,
                                               sheet_name=self.type_df,
                                               skiprows=1)

                        self.colspecs = {}
                        self.dtypes = {}
                        self.colspecs = [(start - 1, end) for start, end in 
                                       zip(self.df_dict['POSIÇÃO INICIAL'],
                                                self.df_dict['POSIÇÃO FINAL'])]
                        self.dtypes  = {}
                        for type_, var in zip(df.TIPO, df.VAR):
                            type_ = type_.strip()
                            dtype = 'string'
                            if tipo == 'C':
                                dtype = 'category'
                            self.dtypes[var] = dtype

                        self.df_dict.columns = ['var', 'name', 'pos', 'end_pos',
                                                'int_part', 'frac_part', 'type']
    def unzip_1960a1991(self):
        with py7zr.SevenZipFile(self.raw_filepath) as zf:
            targets_ = [f for f in zf.getnames() if os.path.splitext(f)[-1] == '.sav'][0]
        dir_path, filename = os.path.split(targets_)
        targets = [dir_path, targets_]
        tmp_path = os.path.join(self.raw_files_path, 'tmp')
        with py7zr.SevenZipFile(self.raw_filepath) as zf:
            zf.extract(path=tmp_path, targets=targets)

        file_path = os.path.join(tmp_path, targets_)
        self.df, self.meta = pyreadstat.read_sav(file_path)
        shutil.rmtree(tmp_path)

    def unzip_2000e2010(self):
        if not hasattr(self, 'database_dict'):
            self.make_database_dict()
        match self.type_db:
            case 'PESS':
                criterion = 'PES'
            case 'DOMI':
                criterion = 'DOM'
        self.df = pd.DataFrame()
        for filepath in self.filepaths:
            with zipfile.ZipFile(filepath, metadata_encoding='cp850') as zf:
                fn = [fn for fn in zf.namelist() if criterion in fn.upper()][0]
                if os.path.splitext(fn)[-1] == '.zip':
                    with zf.open(fn) as f:
                        with zipfile.ZipFile(f, metadata_encoding='cp850') as zf1:
                            fn = [fn for fn in zf1.namelist() if criterion in fn.upper()][0]
                            with zf1.open(fn) as f1:
                                df = pd.read_fwf(f1,
                                                 names=self.df_dict['var'],
                                                 colspecs=self.colspecs,
                                                     dtype=self.dtypes)
                else:
                    with zf.open(fn) as f:
                        df = pd.read_fwf(f,
                                         names=self.df_dict['var'],
                                         colspecs=self.colspecs,
                                             dtype=self.dtypes)

                self.df = pd.concat([self.df, df], ignore_index=True)
        
        if self.uf == 'SP' and self.year == 2010:
            for col, dtype in self.dtypes.items():
                if self.df[col].dtype != dtype:
                    self.df[col] = self.df[col].astype(dtype)

    def unzip(self):
        if not hasattr(self, 'filepaths'):
            self.get_save_raw_database()

        if self.year < 2000:
            self.unzip_1960a1991()
        else:
            self.unzip_2000e2010()

        return self.df

    def str_to_float(self, s, int_part, frac_part):
        if pd.isna(s):
            return pd.NA
        assert(int_part + frac_part == len(s))
        return float(s[:int_part] + '.' + s[int_part:])

    def preprocess_df(self):
        if self.year < 2000: 
            self.preprocess_df_1960a1991()
        else:
            self.preprocess_df_2000e2010()
        return self.df

    def preprocess_df_1960a1991(self):
        if self.year == 1980:
            self.df.drop(columns=['D_R'], inplace=True) #coluna vazia

        for col in self.meta.variable_value_labels.keys():
            self.df[col] = self.df[col].astype('string')
            self.df[col] = self.df[col].astype('category')

        for col in self.df.select_dtypes(exclude='category').column:
            dtype = self.get_min_int_dtype()
            try:
                self.df[col] = self.df[col].abs().astype(dtype)
            except TypeError:
                self.df[col] = self.df[col].abs().astype('Float64')


    def preprocess_df_2000e2010(self):
        if not hasattr(self, 'df'):
            self.unzip()
        float_vars = self.df_dict.loc[self.df_dict.frac_part.notna(),
                                       ['var', 'int_part', 'frac_part', 'size']]
        for var, int_part, frac_part, size in float_vars.itertuples(False, None):
            if self.year == 2000:
                self.df[var]  = self.df[var].str.zfill(size)
            self.df[var] = self.df[var].apply(self.str_to_float,
                                                    args=(int_part, frac_part))
            self.df[var] = self.df[var].astype('Float64')

        for col in self.df.select_dtypes('string').columns:
            if self.year in (2000, 2010) and col == 'V0300':
                continue
            #TODO Refactorization use get_min_int_dtype
            tmp =  pd.to_numeric(self.df[col])
            max_ = tmp.max()
            if max_ >= 2**32:
                dtype = 'UInt64'
            elif max_ >= 2**16:
                dtype = 'UInt32'
            elif max_ >= 2**8:
                dtype = 'UInt16'
            else:
                dtype = 'UInt8'
            self.df[col] = self.df[col].astype(dtype)

        return self.df

    def dict_educacao_superior(self, zf, var, filename, dict_vars, missing_values):
        path = 'Documentação/Anexos Auxiliares'
        pat = re.compile(r'\d{3}')
        with zf.open(os.path.join(path, filename)) as f:
            df = pd.read_excel(f, skiprows=1, dtype='string')
        dict_vars[var] = {key.strip(): value.strip() for key, value
                            in df.iloc[:, :2].dropna().itertuples(index=False, name=None)
                            if pat.search(key)}
        dict_vars[var].update(missing_values)

    def doc2docx(self, zf, path, filename):
            tmp_filepath0 = os.path.join(self.raw_files_path, '~tmp.doc')
            tmp_filepath1 = os.path.join(self.raw_files_path, '~tmp.docx')
            with zf.open(os.path.join(path, filename)) as f:
                with open(tmp_filepath0, 'wb') as f_tmp:
                    f_tmp.write(f.read())
            subprocess.run(['lowriter', '--convert-to', 'docx', tmp_filepath0,
                            '--outdir', self.raw_files_path])
            wordDoc = docx.Document(tmp_filepath1)
            os.remove(tmp_filepath0)
            os.remove(tmp_filepath1)
            return wordDoc

    def make_map_dict(self):
        if self.year < 2000:
            self.map_dict_vars = self.make_map_dict_1960a1991()
        elif self.year == 2000:
            self.map_dict_vars = self.make_map_dict_2000()
        elif self.year == 2010:
            self.map_dict_vars = self.make_map_dict_2010()
        return self.map_dict_vars

    def make_map_dict_1960a1991(self):
        if not hasattr(self, 'meta'):
            with py7zr.SevenZipFile(self.raw_filepath) as zf:
                targets_ = [f for f in zf.getnames() if os.path.splitext(f)[-1] == '.sav'][0]
            dir_path, filename = os.path.split(targets_)
            targets = [dir_path, targets_]
            tmp_path = os.path.join(self.raw_files_path, 'tmp')
            with py7zr.SevenZipFile(self.raw_filepath) as zf:
                zf.extract(path=tmp_path, targets=targets)

            file_path = os.path.join(tmp_path, targets_)
            _, self.meta = pyreadstat.read_sav(file_path, metadataonly=True)
            shutil.rmtree(tmp_path)

        
        map_var_dict = {}
        for key in self.meta.column_names:
            dict_map = self.meta.variable_value_labels.get(key, pd.NA)
            if pd.notna(dict_map):
                map_var_dict[key] = {str(int(k)): v for k, v in dict_map.items()}

        if self.year == 1970:
            #correção de um erro na codificação do SPSS
            #variável de V035 alfabetização e V036 frequenta a escola
            map_var_dict['V035'] = {'0': 'Sem declaração',
                                    '1': 'Sim',
                                    '2': 'Não'}
            map_var_dict['V036'] = {'0': 'Sem declaração',
                                    '1': 'Sim',
                                    '2': 'Não'}
        elif self.year == 1980:
            #Algumas variáveis não estão codificadas no arquivo .sav
            map_var_dict['V517'] = {'0': 'menos de 1 ano',
                                    '1': '1 ano',
                                    '2': '2 anos',
                                    '3': '3 anos',
                                    '4': '4 anos',
                                    '5': '5 anos',
                                    '6': '6 a 9 anos',
                                    '7': '7 a 10 anos',
                                    '8': 'nasceu',
                                    '9': 'sem declaração'},

            map_var_dict['V521'] = {'0': 'nenhuma',
                                    '1': 'primário',
                                    '2': 'ginasial médio',
                                    '3': '1º grau',
                                    '4': '2º grau',
                                    '5': 'colegial médio',
                                    '6': 'supletivo 1º grau',
                                    '7': 'supletivo 2º grau',
                                    '8': 'superior',
                                    '9': 'sem declaração'}
            #TODO há outras


        df = pd.DataFrame({'COD_VAR': [key for key in self.meta.column_names],
                           'NOME_VAR': [self.meta.column_names_to_labels.get(key, pd.NA) 
                                        for key in self.meta.column_names],
                           'MAP_VAR': [map_var_dict.get(key, pd.NA)
                                       for key in self.meta.column_names]})
        df.to_excel(f'{self.path_dict}.xlsx', index=False)
        df.to_pickle(f'{self.path_dict}.pickle')
        return df

    def make_map_dict_2000(self):
        path = 'Arquivos Auxiliares'
        external_vars = dict() 
        docpath = glob.glob(f'{self.raw_files_path}/*[Dd]oc*.zip')[0]
        with zipfile.ZipFile(docpath, metadata_encoding='cp850') as zf:

            #V4250 = Municípios
            with zf.open(os.path.join(path, 'Municipios-V4250.xls')) as f:
                df_ = pd.read_excel(f, dtype='string')
            external_vars['V4250'] = {key.strip(): value.strip() for key, value
                                      in df_.dropna().itertuples(False, None)}

            #V4276 = Municípios e países estrangeiros
            with zf.open(os.path.join(path, 'Municipios e Pais Estrangeiro - V4276.xls')) as f:
                df_ = pd.read_excel(f, dtype='string')
            external_vars['V4276'] = {key.strip(): value.strip() for key, value
                                      in df_.dropna().itertuples(False, None)}

            #V4279 = Países estrangeiros
            with zf.open(os.path.join(path, 'Estrutura ONU V4279.xls')) as f:
                df_ = pd.read_excel(f, dtype='string', skiprows=3, na_values=[' '])
            external_vars['V4279'] = {key.strip(): value.strip() for value, key
                                      in df_.dropna(subset='CODIGO').itertuples(False, None)}

            #V4239 = Países estrangeiros
            with zf.open(os.path.join(path, 'Estrutura ONU V4239.xls')) as f:
                df_ = pd.read_excel(f, dtype='string', skiprows=3, na_values=[' '])
            external_vars['V4239'] = {key.strip(): value.strip() for value, key
                                      in df_.dropna(subset='CODIGO').itertuples(False, None)}

            #V4219 e V4269 = Países estrangeiros e UFs
            with zf.open(os.path.join(path, 'Estrutura ONU V4219, V4269.xls')) as f:
                df_ = pd.read_excel(f, dtype='string', skiprows=3, na_values=[' '])
            var_ext_tmp = {key.strip(): value.strip() for value, key
                           in df_.dropna(subset='CODIGO').itertuples(False, None)
                           if key.isdigit()}
            external_vars['V4219'] = var_ext_tmp 
            external_vars['V4269'] = var_ext_tmp 

            #V4230 = Países estrangeiros e UFs
            with zf.open(os.path.join(path, 'Estrutura Migracao V4230.xls')) as f:
                df_ = pd.read_excel(f, dtype='string', skiprows=2, na_values=[' '])
            external_vars['V4230'] = {key.strip(): value.strip() for key, value
                                      in df_.dropna(subset='CODIGOS').itertuples(False, None)
                                      if key.isdigit()}

            #V4210 e V4260 = Países estrangeiros e UFs
            with zf.open(os.path.join(path, 'Estrutura Migracao V4210, V4260.xls')) as f:
                df_ = pd.read_excel(f, dtype='string', skiprows=2, na_values=[' '])
            var_ext_tmp = {key.strip(): value.strip() for key, value
                           in df_.dropna(subset='CODIGO').itertuples(False, None)
                           if key.isdigit()}
            external_vars['V4210'] = var_ext_tmp
            external_vars['V4260'] = var_ext_tmp

            #V4355 = Cursos Superiores; e area_de_conhecimento
            with zf.open(os.path.join(path, 'Cursos Superiores - Estrutura V4535.xls')) as f: #Houve algum erro de digitação, pois a variável correta é V4355, apesar de o arquivo se referir à variável V4535, a documentação também se refere ao documento com o nome da variável errado.
                df_ = pd.read_excel(f, dtype='string', skiprows=5, na_values=[' '])
            external_vars['V4355'] = {key.strip(): value.strip() for key, value
                                      in df_.iloc[:, 1:].dropna(subset='Código').itertuples(False, None)
                                      if key.isdigit()}
            external_vars['V4355']['02'] = 'Não Superior'
            external_vars['cursos_superiores_area_de_conhecimento'] = {}
            areas = []
            new_area = None
            for line in df_.iloc[:, 0].dropna():
                if line[0].isdigit():
                    if new_area is not None:
                        areas.append(new_area.strip())
                    new_area = line
                else:
                    new_area += line
            areas.append(new_area)
            for area in areas:
                key, value = area.split('-')
                for k in re.findall(r'\d', key):
                    external_vars['cursos_superiores_area_de_conhecimento'][k] = value.strip()

            #V4354 = Cursos Superiores; areas_especificas e areas_gerais
            with zf.open(os.path.join(path, 'Cursos Superiores - Estrutura V4534.xls')) as f: #Mesmo caso do erro da V4355
                df_ = pd.read_excel(f, dtype='string', skiprows=4, na_values=[' '])
            external_vars['V4354'] = {}
            external_vars['cursos_superiores_areas_especificas'] = {}
            external_vars['cursos_superiores_areas_gerais'] = {}
            for line in df_.iloc[:, 2].dropna():
                key, value = line.split(maxsplit=1)
                external_vars['V4354'][key.strip()] = value.strip()
            for line in df_.iloc[:, 1].dropna():
                key, value = line.split(maxsplit=1)
                external_vars['cursos_superiores_areas_especificas'][key.strip()] = value.strip()
            for line in df_.iloc[:, 0].dropna():
                key, value = line.split(maxsplit=1)
                external_vars['cursos_superiores_areas_gerais'][key.strip()] = value.strip()

            #V1004 = Região Metropolitana
            external_vars['V1004'] = {}
            with zf.open(os.path.join(path, 'V1004.txt')) as f:
                for line in f.readlines():
                    line = line.decode('windows-1252')
                    if line[0].isdigit():
                        key, value = line.split('-', 1)
                        external_vars['V1004'][key.strip()] = value.strip()

            #V4451 = Código antigo da ocupação, relativo a 91
            with zf.open(os.path.join(path, 'Ocupacao91-Estrutura.xls')) as f:
                df_ = pd.read_excel(f, dtype='string', skiprows=2, header=None, names=['key', 'value'], na_values=[' '])
            external_vars['V4451'] = {key.strip(): value.strip() for key, value
                                      in df_.dropna(subset='key').itertuples(False, None)
                                      if key.isdigit()}
            external_vars['V4451']['927'] = 'OUTRAS OCUPACOES OU OCUPACOES MAL DEFINIDAS'
            external_vars['V4451']['000'] = 'OUTRAS OCUPACOES OU OCUPACOES MAL DEFINIDAS'

            df_.loc[df_['key'].notna(), 'value'] = pd.NA
            df_['value'] = df_['value'].ffill()
            external_vars['ocupacao_91_gg'] = {key.strip(): value.strip() for key, value
                                               in df_.dropna(subset='key').itertuples(False, None)
                                               if key.isdigit()}
            external_vars['ocupacao_91_gg']['927'] = 'OUTRAS OCUPACOES OU OCUPACOES MAL DEFINIDAS'
            external_vars['ocupacao_91_gg']['000'] = 'OUTRAS OCUPACOES OU OCUPACOES MAL DEFINIDAS'
            external_vars['ocupacao_91_gg']['999'] = 'OUTRAS OCUPACOES OU OCUPACOES MAL DEFINIDAS'


            #V4452 = Código novo da ocupação, relativo a 2000
            wordDoc = self.doc2docx(zf, path, 'Ocupacao-Estrutura.doc')
            rows_ = []
            for i, table in enumerate(wordDoc.tables):
                for row in table.rows:
                    row_ = []
                    for j, cell in enumerate(row.cells):
                        if i >= 1 and j == 0 and cell.text.startswith(('C', 'G')):
                            break
                        row_.append(cell.text)
                    if row_:
                        rows_.append(row_)
            df_ = pd.DataFrame(rows_[2:], columns=rows_[1])
            external_vars['V4452'] = {key.strip(): value.strip() for key, value
                                      in df_[['Grupo de base', 'Titulação']]
                                            .dropna(subset='Grupo de base')
                                            .itertuples(False, None)
                                      if key.isdigit()}
            external_vars['ocupacao_2000_sg'] = {key.strip()[:3]: value.strip() for key, value
                                      in df_[['Subgrupo', 'Titulação']]
                                            .dropna(subset='Subgrupo')
                                            .itertuples(False, None)
                                      if key.isdigit()}
            external_vars['ocupacao_2000_sgp'] = {key.strip()[:2]: value.strip() for key, value
                                      in df_[['Subgrupo principal', 'Titulação']]
                                            .dropna(subset='Subgrupo principal')
                                            .itertuples(False, None)
                                      if key.isdigit()}
            external_vars['ocupacao_2000_gg'] = {key.strip()[:1]: value.strip() for key, value
                                      in df_[['Grande Grupo', 'Titulação']]
                                            .dropna(subset='Grande Grupo')
                                            .itertuples(False, None)
                                      if key.isdigit()}
            external_vars['V4452']['0000'] = 'OCUPAÇÕES MAL ESPECIFICADAS'
            external_vars['ocupacao_2000_sg']['000'] = 'OCUPAÇÕES MAL ESPECIFICADAS' 
            external_vars['ocupacao_2000_sgp']['00'] = 'OCUPAÇÕES MAL ESPECIFICADAS'
            external_vars['ocupacao_2000_gg']['0'] = 'OCUPAÇÕES MAL ESPECIFICADAS' 

            #V4090 = Estrutura de Religião e grandes grupos de religião
            wordDoc = self.doc2docx(zf, path, 'Estrutura de Religiao - V4090.doc')
            rows_ = []
            for table in wordDoc.tables:
                for row in table.rows:
                    row_ = []
                    for j, cell in enumerate(row.cells):
                        if (j == 0 and cell.text.startswith('R')) or j >= 2 or cell.text == '':
                            break
                        row_.append(cell.text)
                    if row_:
                        rows_.append(row_)
            df_ = pd.DataFrame(rows_, columns=['key', 'value'])
            external_vars['V4090'] = {key.strip(): value.strip() for key, value
                                      in df_.itertuples(False, None)
                                      if key.isdigit()}
            external_vars['religiao_gg'] = {line.split(maxsplit=1)[0].strip(): line.split(maxsplit=1)[1].strip() for line
                                            in df_['key'] if not line.isdigit()}
            external_vars['religiao_gg']['00'] = 'SEM RELIGIÂO'
            external_vars['religiao_gg']['99'] = 'SEM DECLARAÇÃO'

            #V4461 = Código antigo da atividade de trabalho relativo a 91
            with zf.open(os.path.join(path, 'Atividade91-Estrutura.xls')) as f:
                df_ = pd.read_excel(f, dtype='string', header=None, names=['key', 'value'], na_values=[' '])
            external_vars['V4461'] = {key.strip(): value.strip() for key, value
                                      in df_.dropna(subset='key').itertuples(False, None)
                                      if key.isdigit()}

            df_.loc[df_['key'].notna(), 'value'] = pd.NA
            df_['value'] = df_['value'].ffill()
            external_vars['atividade_91_gg'] = {key.strip(): value.strip() for key, value
                                                in df_.dropna(subset='key').itertuples(False, None)
                                                if key.isdigit()}

            #V4462 = Código novo da atividade
            with zf.open(os.path.join(path, 'CnaeDom-Estrutura.xls')) as f:
                df_ = pd.read_excel(f, dtype='string', skiprows=2, header=None, names=['key', 'value'], na_values=[' '])
            external_vars['V4462'] = {key.strip(): value.strip() for key, value
                                      in df_.dropna(subset='key').itertuples(False, None)
                                      if re.match(r'\d{5}', key)}
            external_vars['atividade-sgp'] = {key.strip(): value.strip() for key, value
                                              in df_.dropna(subset='key').itertuples(False, None)
                                              if re.match(r'\b\d{2}\b', key)}
            df_.loc[df_['key'].notna(), 'value'] = pd.NA
            df_['value'] = df_['value'].ffill()
            external_vars['atividade-gg'] = {key.strip(): value.split('-', 1)[-1].strip() for key, value
                                             in df_.dropna(subset='key').itertuples(False, None)
                                             if re.match(r'\d{5}', key)}

            #Divisão territorial brasileira
            with zf.open(os.path.join(path, 'Divisao Territorial Brasileira.xls')) as f:
                dfs_ = pd.read_excel(f, dtype='string', skiprows=1, header=None, names=['key', 'value'], na_values=[' '], 
                                     sheet_name=['Mesorregião', 'Microrregião', 'Município', 'Distrito', 'Subdistrito'])

            #V1002 = Mesorregião; V1003 = Microrregião; V0103 e V1103 = Municipio; V0104 = Distrito; V0105 = Subdistrito
            vars_ = {'V1002': 'Mesorregião',
                     'V1003': 'Microrregião',
                     'V0103': 'Município',
                     'V1103': 'Município',
                     'V0104': 'Distrito', 
                     'V0105': 'Subdistrito'}
            for k, v in vars_.items():
                external_vars[k] = {key.strip(): value.strip() for key, value
                                    in dfs_[v].dropna(subset='key').itertuples(False, None)
                                    if key.isdigit()}

            wordDoc = self.doc2docx(zf, 'Documentacao', 'Documentação.doc')
            rows_ = []
            for table in wordDoc.tables:
                for row in table.rows:
                    row_ = []
                    for j, cell in enumerate(row.cells):
                        if (j == 0 and (cell.text.startswith('VA') or not cell.text.startswith(('V', 'M')))) or j >= 2:
                            break
                        if j == 0:
                            row_.append(cell.text)
                        else:
                            var_dict = {}
                            for i, e in enumerate(cell.text.split('\n')):
                                if i == 0:
                                    row_.append(e)
                                elif e[0].isdigit() and e.find('-') != -1:
                                    k, v = e.split('-', 1)
                                    var_dict[k.strip()] = v.strip()
                            row_.append(var_dict)
                    if row_:
                        rows_.append(row_)
            df = pd.DataFrame(rows_, columns=['COD_VAR', 'NOME_VAR', 'MAP_VAR'])

        #Ajustes manuais
        df.loc[df.COD_VAR == 'V0104', 'NOME_VAR'] = 'CÓDIGO DO DISTRITO'
        df.loc[df.COD_VAR == 'V0300', 'NOME_VAR'] = 'IDENTIFICAÇÃO DO DOMICÍLIO'
        df.loc[df.COD_VAR == 'V0400', 'NOME_VAR'] = 'NÚMERO DE ORDEM DA PESSOA RECENSEADA'

        df = df.drop_duplicates(subset=['COD_VAR', 'NOME_VAR'])

        extra_vars = []
        for k, v in external_vars.items():
            filter_ = df.COD_VAR == k
            if filter_.sum():
                df.loc[df.COD_VAR == k, 'MAP_VAR'] = [v] #Hack(?) para o pandas aceitar um elemento de tipo dicionário
            else:
                extra_vars.append([k, k, v])

        df_extra = pd.DataFrame(extra_vars, columns=['COD_VAR', 'NOME_VAR', 'MAP_VAR'])
        df = pd.concat([df, df_extra], ignore_index=True)
        
        df.to_excel(f'{self.path_dict}.xlsx', index=False)
        df.to_pickle(f'{self.path_dict}.pickle')

        #TODO indicar nas variáveis extras o quanto se deve fazer slice (self.df[var].str.slice(stop=x).astype('category')
        #TODO indicar também naquelas em que não é possível realizar o slice e aplicar o map antes do agrupamento na função crosstable

        return df


    def make_map_dict_2010(self):
        path = 'Documentação/Anexos Auxiliares'
        path_regioes = 'Documentação/Divisão Territorial do Brasil/'
        external_vars = collections.defaultdict(dict)
        docpath = glob.glob(f'{self.raw_files_path}/*[Dd]oc*.zip')[0]
        with zipfile.ZipFile(docpath, metadata_encoding='cp850') as zf:
            with zf.open(os.path.join(path, 'Atividade CNAE_DOM 2.0 2010.xls')) as f:
                df_ = pd.read_excel(f, skiprows=1, dtype='string')
            external_vars['V6471'] = {key.strip(): value.strip() for key, value
                                      in df_.iloc[:, 2:4]
                                            .dropna()
                                            .itertuples(index=False, name=None)}

            self.dict_educacao_superior(zf,
                                        'V6356',
                                        'Cursos Doutorado_Estrutura 2010.xls',
                                        external_vars,
                                        {'097': 'NÃO SABE E DOUTORADO NÃO ESPECIFICADO'})
            self.dict_educacao_superior(zf,
                                        'V6354',
                                        'Cursos Mestrado_Estrutura 2010.xls',
                                        external_vars,
                                        {'097': 'NÃO SABE E MESTRADO NÃO ESPECIFICADO'})
            self.dict_educacao_superior(zf,
                                        'V6352',
                                        'Cursos Superiores_Estrutura 2010.xls',
                                        external_vars,
                                        {'085': 'NÃO SABE E SUPERIOR NÃO ESPECIFICADO'})

            pat = re.compile(r'\d{4}')
            with zf.open(os.path.join(path, 'Ocupação COD 2010.xls')) as f:
                df_ = pd.read_excel(f, skiprows=1, dtype='string')
            external_vars['V6461'] = {key: value for key, value
                                      in df_.iloc[:, 0:2]
                                            .dropna()
                                            .itertuples(index=False, name=None)
                                      if pat.search(key)}

            with zf.open(os.path.join(path, 'Migração e deslocamento _Unidades da Federação.xls')) as f:
                df_ = pd.read_excel(f, skiprows=5, dtype='string')
            dict_mig_uf = {key.strip(): value.strip() for value, key
                           in df_.dropna().itertuples(index=False, name=None)}
            for var in ('V6222', 'V6252', 'V6262', 'V6362', 'V6602'):
                external_vars[var] = dict_mig_uf

            with zf.open(os.path.join(path, 'Migração e deslocamento _Municípios.xls')) as f:
                df_ = pd.read_excel(f, skiprows=6, dtype='string')
            dict_mig_mun = {key.strip(): value.strip() for value, key
                           in df_.iloc[:, 1:3]
                                  .dropna()
                                  .itertuples(index=False, name=None)}
            for var in ('V6254', 'V6264', 'V6364', 'V6604'):
                external_vars[var] = dict_mig_mun

            with zf.open(os.path.join(path, 'Migração e Deslocamento_Paises estrangeiros.xls')) as f:
                df_ = pd.read_excel(f, skiprows=7, dtype='string')
            dict_mig_pais = {key.strip(): value.strip() for value, key
                           in df_.iloc[:, 1:3]
                                   .dropna()
                                   .itertuples(index=False, name=None)}
            for var in ('V3061', 'V6224', 'V6256', 'V6266', 'V6366', 'V6606'):
                external_vars[var] = dict_mig_pais

            pat = re.compile(r'\d{3}')
            with zf.open(os.path.join(path, 'Religião 2010.txt')) as f:
                for line in f.readlines():
                    line = line.decode('latin_1').strip()
                    if pat.search(line):
                        key, value = line.split(maxsplit=1)
                        external_vars['V6121'][key] = value

            with zf.open(os.path.join(path, 'Estrutura atividade CD2000.xls')) as f:
                df_ = pd.read_excel(f, dtype='string')

            pat = re.compile(r'\d{5}')
            dict_ativ_2000 = {key.strip(): value.strip() for key, value
                                 in df_.dropna().itertuples(index=False, name=None)
                                 if pat.search(key)}
            external_vars['V6472'] = dict_ativ_2000

            with zf.open(os.path.join(path_regioes,
                                      'Unidades da Federação, Mesorregiões, microrregiões e municípios 2010.xls')) as f:
                df_ = pd.read_excel(f, skiprows=2, dtype='string')
                for _, _, meso, nome_meso, micro, nome_micro, mun, nome_mun in df_.itertuples(index=False, name=None):
                    external_vars['V1002'][meso.strip()] = nome_meso.strip()
                    external_vars['V1003'][micro.strip()] = nome_micro.strip()
                    external_vars['V0002'][mun.strip()] = nome_mun.strip()

        wordDoc = docx.Document(os.path.join(self.raw_files_path, 'estrutura ocupacao CD2000.docx'))
        for table in wordDoc.tables:
            rows_ = []
            for row in table.rows:
                row_ = []
                for cell in row.cells:
                    row_.append(cell.text)
                rows_.append(row_)
        dict_ocu_2000 = {key: value for key, value
                         in pd.DataFrame(data=rows_[2:], columns=rows_[1])
                         .iloc[:, 3:5].itertuples(index=False, name=None) if key}
        dict_ocu_2000['0000']= 'OCUPAÇÕES MAL ESPECIFICADAS'
        external_vars['V6462'] = dict_ocu_2000

        cod_vars_dict = collections.defaultdict(dict)
        pat = re.compile(r'\d+\s*-')
        names = {}
        if not hasattr(self, 'df_dict'):
            self.make_database_dict()
        for var, e in zip(self.df_dict[self.type_db].VAR,
                          self.df_dict[self.type_db].NOME):
            flag = False
            for i, line in enumerate(e.split('\n')):
                if not i:
                    names[var] = line.strip(' :')
                if pat.search(line):
                    key, value = [n.strip() for n in line.split('-', 1)]
                    if key not in cod_vars_dict[var]:
                        cod_vars_dict[var][key] = value
                    flag = True

        missing_vars = {
            'V5110': {'1': 'Contribuintes',
                      '2': 'Não contribuintes'},
            'V5120': {'1': 'Contribuintes',
                      '2': 'Não contribuintes'},
        }

        cod_vars_dict.update(missing_vars)
        cod_vars_dict.update(external_vars)

        df = pd.DataFrame({'COD_VAR': [key for key in names],
              'NOME_VAR': [nome for nome in names.values()],
              'MAP_VAR': [cod_vars_dict.get(key, pd.NA) for key in names]})
        df.to_excel(f'{self.path_dict}.xlsx', index=False)
        df.to_pickle(f'{self.path_dict}.pickle')
        return df

    def get_coded_var(self, var):
        if self.year == 2010:
            if var == 'V0002':
                col = self.cod_mun
            elif var == 'V1002':
                col = self.cod_meso
            elif var == 'V1003':
                col = self.cod_micro
            else:
                col = self.df[var]
        else:
            col = self.df[var]
        return col.map(self.get_map_var(var)[1])

    def get_df(self, filetype='parquet', **kwargs):
        if filetype not in SUPPORTED_FTs:
            raise ValueError

        if self.uf == 'ALL' and self.year >= 2000:
            self.dir_path = os.path.join(self.path, FILETYPES_PATH[filetype])
            if not os.path.isdir(self.dir_path):
                os.mkdir(self.dir_path)
                
            self.dest_filepath = os.path.join(self.dir_path,
                                              f'{self.filename}.{filetype}')
            if os.path.isfile(self.dest_filepath):
                print_info(f'Arquivo {self.dest_filepath} já existente')
                read_fun = getattr(pd, f'read_{filetype}')
                self.df = read_fun(self.dest_filepath, **kwargs)

            else:
                all_ufs = [f for f in glob.glob(f'{self.dir_path}/*.{filetype}')
                              if 'ALL' not in f]
                is_complete = True
                for uf in UF_SIGLA_NOME:
                    has_processed = False
                    for f in all_ufs:
                        if uf in f:
                            has_processed = True
                    if not has_processed:
                        print_error(f'A uf {uf} não foi ainda processada')
                        is_complete = False
                if not is_complete:
                    print_error('É preciso processar todas as ufs antes de juntá-las todas')
                    raise ValueError
                print_info('Todas as ufs já foram processadas, preparando para juntá-las')
                self.df = pd.DataFrame()
                for f in all_ufs:
                    print_info(f'Anexando o arquivo {f} no DataFrame')
                    df_tmp = pd.read_parquet(f)
                    self.df = pd.concat([self.df, df_tmp], ignore_index=True)

                for col in self.df.select_dtypes(object).columns:
                    if col == 'V0300':
                        continue
                    self.df[col] = self.df[col].astype('category')

                self.save(filetype=self.filetype)

        else:
            self.df = super().get_df(filetype, **kwargs)

        if self.year == 2010:
            self.cod_mun = (self.df.V0001.astype('string')
                            + self.df.V0002.astype('string')).astype('category')
            self.cod_meso = (self.df.V0001.astype('string')
                            + self.df.V1002.astype('string')).astype('category')
            self.cod_micro = (self.df.V0001.astype('string')
                            + self.df.V1003.astype('string')).astype('category')
        self.complementary_vars()
        return self.df

    def complementary_vars(self):
        match self.year:
            case 1960:
                educacao_1960(self.df)
                # idade_1960(self.df)
            case 1970:
                educacao_1970(self.df)
                # idade_1970(self.df)
            case 1980:
                educacao_1980(self.df)
                # idade_1980(self.df)
            case 1991:
                educacao_1991(self.df)
                # idade_1991(self.df)
            case 2000:
                educacao_2000(self.df)
                # idade_2000(self.df)
            case 2010:
                idade_2010(self.df)
                educacao_2010(self.df)


NAO_CONCLUIU_SEM_DECLARACAO = 'Não concluiu e sem declaração de alfabetização'
NAO_CONCLUIU_ANALF = 'Não concluiu e não alfabetizado'
NAO_CONCLUIU_ALFA = 'Não concluiu e alfabetizado'
EF_AI = 'Ensino Fundamental - Anos Iniciais'
EF_AF = 'Ensino Fundamental - Anos Finais'
EM = 'Ensino Médio'
ES = 'Educação Superior'

E_GRADES = [
    NAO_CONCLUIU_SEM_DECLARACAO,
    NAO_CONCLUIU_ANALF,
    NAO_CONCLUIU_ALFA,
    EF_AI,
    EF_AF,
    EM,
    ES,
]

NAO_FREQUENTA = 'Não frequenta'
EF_SD = 'Ensino Fundamental - Sem declaração de série'
EF_0  = 'Ensino Fundamental - 1º Ano'
EF_1  = 'Ensino Fundamental - 1ª Série'
EF_2  = 'Ensino Fundamental - 2ª Série'
EF_3  = 'Ensino Fundamental - 3ª Série'
EF_4  = 'Ensino Fundamental - 4ª Série'
EF_5  = 'Ensino Fundamental - 5ª Série'
EF_6  = 'Ensino Fundamental - 6ª Série'
EF_7  = 'Ensino Fundamental - 7ª Série'
EF_8  = 'Ensino Fundamental - 8ª Série'
EM_SD = 'Ensino Médio - Sem declaração de série'
EM_1  = 'Ensino Médio - 1ª Série'  
EM_2  = 'Ensino Médio - 2ª Série'
EM_3  = 'Ensino Médio - 3ª Série'
ES_SD = 'Ensino Superior - Sem declaração de série'
ES_1  = 'Ensino Superior - 1º Ano'
ES_2  = 'Ensino Superior - 2º Ano'
ES_3  = 'Ensino Superior - 3º Ano'
ES_4  = 'Ensino Superior - 4º Ano'
ES_5  = 'Ensino Superior - 5º Ano'
ES_6  = 'Ensino Superior - 6º Ano'
PE    = 'Pré-escolar'
AA    = 'Alfabetização para adultos'
S1    = 'Supletivo 1º Grau'
S2    = 'Supletivo 2º Grau'
V     = 'Vestibular'
PG    = 'Pós-graduação'

F_GRADES = [
    NAO_FREQUENTA,
    PE, AA, EF_SD,
    EF_0, EF_1, EF_2, EF_3, EF_4,
    EF_5, EF_6, EF_7, EF_8,
    S1, EM_SD,
    EM_1, EM_2, EM_3,
    S2, V, ES_SD,
    ES_1, ES_2, ES_3,
    ES_4, ES_5, ES_6,
    PG,
]

CAT_E_TYPES = CategoricalDtype(categories=E_GRADES, ordered=True)
CAT_F_TYPES = CategoricalDtype(categories=F_GRADES, ordered=True)

MAX_EF_AI = 4
MAX_EF_AF = 8
MAX_EM    = 11
MAX_ES    = 15

C_ANOS_ESC   = 'anos_esc'
C_ETAPA_CONC = 'etapa_concluida'
C_FREQ = 'etapa_frequentada'

#As funções de harmonização das informações educacionais são, em parte, uma conversão dos scripts
# em R elaborados por @antrologos, disponível em https://github.com/antrologos/VariaveisHarmonizadasDataCEM/
# e https://github.com/antrologos/harmonizeIBGE/
def educacao_1960(df):
    '''
    V211 - Alfabetização
    ====================
    0   Lê e Freqüenta Escola
    1   Lê e não Freqüenta Escola
    2   não Lê e Freqüenta Escola
    3   não Lê e não Freqüenta Escola
    4   Ignorada
        Não aplicável (4 anos de idade ou menos) ou Informação Faltante (Registro Corrompido)

    V212 - Última série concluída
    =============================
    4    Primeira Série
    5    Segunda Série
    6    Terceira Série
    7    Quarta Série
    8    Quinta Série
    9    Sexta Série
    0    Esta cursando o Primeiro ano do Elementar (não possui série concluída)
    1    Nunca Frequentou Escola
    2    Ignorado
         Não aplicável (4 anos de idade ou menos) ou Informação Faltante (Registro Corrompido)

    V213 - Grau do curso
    ====================
    2    Elementar
    3    Médio Primeiro Ciclo
    4    Médio Segundo Ciclo
    5    Superior
    6    Ignorado
    1    Nunca Frequentou Escola
    0    Esta cursando o Primeiro ano do Elementar (não possui série concluída)
         Não aplicável (4 anos de idade ou menos) ou Informação Faltante (Registro Corrompido)

    V214 - Curso completo 
    =====================
    0    Sem Curso Completo
    10   Primário/Elementar
    11   Comercial - Elementar
    13   Saúde e Serviços Sanitários - Elementar
    14   Militar Elementar - Elementar
    15   Agricultura e Pecuária - Elementar
    16   Emendativo - Elementar
    17   Industrial - Elementar
    19   Outros - Elementar
    20   Ginasial
    21   Comercial - 1º Grau / Médio 1º Ciclo
    22   Normal/Pedagógico - 1º Grau / Médio 1º Ciclo
    24   Militar - 1º Grau / Médio 1º Ciclo
    25   Agricultura e Pecuária - 1º Grau / Médio 1º Ciclo
    26   Emendativo - 1º Grau / Médio 1º Ciclo
    27   Industrial - 1º Grau / Médio 1º Ciclo
    29   Outros - 1º Grau / Médio 1º Ciclo
    33   Serviços Sanitários - 1º Grau / Médio 1º Ciclo
    34   Militar Médio - 1º Grau / Médio 1º Ciclo
    36   Educação Física - 1º Grau / Médio 1º Ciclo
    38   Eclesiástico - 1º Grau / Médio 1º Ciclo
    39   Outros Níveis Médios (1º Grau / Médio 1º Ciclo)
    40   Colegial/Científico
    41   Comercial - 2º Grau
    42   Normal/Pedagógico - 2º Grau
    44   Militar - 2º Grau
    45   Agricultura e Pecuária - 2º Grau
    47   Industrial - 2º Grau
    49   Outros - 2º Grau
    50   Geografia e História - Superior
    51   História Natural - Superior
    52   Letras - Superior
    53   Matemática, Física, Química, Desenho - Superior
    54   Outros Cursos Superiores (Pedagogia, Filosofia, Ciências Sociais, Teologia)
    57   Belas Artes - Superior
    60   Medicina - Superior
    61   Farmácia - Superior
    62   Odontologia - Superior
    63   Veterinária - Superior
    64   Engenharia - Superior
    65   Arquitetura - Superior
    66   Química Industrial - Superior
    67   Direito - Superior
    68   Agronomia - Superior
    70   Ciências Econômicas - Superior
    71   Estatística  - Superior
    72   Artes Domésticas  - Superior
    73   Saúde, Enfermagem e Serviços Sanitários - Superior
    74   Militar - Superior
    76   Educação Física - Superior
    78   Eclesiástico - Superior (Teologia e Filosofia para formação eclesiástica)
    79   Outros - Nível Superior (Adm. Pública, Música, Jornalismo, Museologia etc)
    89   Curso com grau não especificado
    99   Ignorado 
         Não aplicável (9 anos de idade ou menos) ou Informação Faltante (Registro Corrompido)
    '''
    
    #Anos base de estudo para cada grau
    yearsStage = {
        '2': 0,
        '3': 4,
        '4': 8,
        '5': 11
    }
    #Anos de estuda para as séries
    yearsSeries = {
        '4': 1,
        '5': 2,
        '6': 3,
        '7': 4,
        '8': 5,
        '9': 6
    }
    
    #Coluna de anos de escolaridade
    df[C_ANOS_ESC] = df.V213.map(yearsStage) + df.V212.map(yearsSeries)
    df.loc[df.V213.isin({'0', '1'}), C_ANOS_ESC] = 0
    #Aplicar teto para os graus de escolaridade
    #Elementar: 4 anos; Médio 1º Ciclo: 8; Médio 2º Ciclo: 11; Superior: 15
    df.loc[(df[C_ANOS_ESC] > MAX_EF_AI) & (df.V213 == '2'), C_ANOS_ESC] = MAX_EF_AI
    df.loc[(df[C_ANOS_ESC] > MAX_EF_AF) & (df.V213 == '3'), C_ANOS_ESC] = MAX_EF_AF
    df.loc[(df[C_ANOS_ESC] > MAX_EM)    & (df.V213 == '4'), C_ANOS_ESC] = MAX_EM
    df.loc[(df[C_ANOS_ESC] > MAX_ES)    & (df.V213 == '5'), C_ANOS_ESC] = MAX_ES
    df[C_ANOS_ESC] = df[C_ANOS_ESC].astype('UInt8')

    #Coluna de conclusão de etapa

    #Preparar algumas variáveis
    V214i = df.V214.astype('UInt16')
    filter_no_grade = (V214i == 0) | (df.V212 == '1') | (df.V213 == '1')
    df[C_ETAPA_CONC] = pd.NA

    df.loc[filter_no_grade & df.V211.isin({'2', '3'}), C_ETAPA_CONC] = NAO_CONCLUIU_ANALF
    df.loc[filter_no_grade & df.V211.isin({'0', '1'}), C_ETAPA_CONC] = NAO_CONCLUIU_ALFA
    df.loc[V214i.between(10, 19), C_ETAPA_CONC] = EF_AI
    df.loc[V214i.between(20, 29), C_ETAPA_CONC] = EF_AF
    df.loc[V214i.between(30, 49), C_ETAPA_CONC] = EM
    df.loc[V214i.between(50, 79), C_ETAPA_CONC] = ES
    df[C_ETAPA_CONC] = df[C_ETAPA_CONC].astype(CAT_E_TYPES)

def educacao_1970(df):
    '''
    V035 - Alfabetização
    ====================
    0     Sem declaração
    1     Sim
    2     Não
    --------------------
    Nota: O arquivo .sav disponibilizado pelo CEM não segue os valores
    da documentação, por essa razão, realizei essa compatibilização.

    V036 - Frequenta a escola
    ========================= 
    0     Sem declaração
    1     Sim
    2     Não
    -------------------------
    Nota: ver nota acima

    V037 - Última série que concluiu com aprovação neste curso que frequentou anteriormente
    ===================================================================================== 
    1     1ª série do elementar
    2     1ª série
    3     2ª série
    4     3ª série
    5     4ª série
    6     5ª ou 6ª série
    7     Admissão ou vestibular
    8     Artigo 99
    9     Alfabetização de adultos
    0     Nenhuma ou sem declaração
    ------------------------------------------------------------------------------------- 
    Nota: O Art. 99 da LDB/1961: "Aos maiores de dezesseis anos será permitida a obtenção
    de certificados de conclusão do curso ginasial, mediante a prestação de exames de
    madureza, após estudos realizados sem observância do regime escolar.
    Parágrafo único. Nas mesmas condições, permitir-se-á a obtenção do certificado de
    conclusão do curso colegial aos maiores de dezenove anos."
    (Redação dada pelo Decreto-Lei nº 709, 1969) 

    V038 - Último grau concluído com aprovação
    ==========================================
    1     Primário/elementar
    2     Ginasial/médio 1º ciclo
    3     Colegial/médio 2º ciclo
    4     Superior
    5     Nunca frequentou escola
    0     Sem declaração
    ------------------------------------------

    V039 - Espécie de curso concluído
    =================================
    0     Sem declaração
    10    Primário
    11    Agrícola elementar
    12    Comercial elementar
    19    Industrial elementar
    21    Militar elementar
    22    Normal elementar
    27    Outros elementar
    28    Emendativo elementar
    30    Ginasial
    31    Agrícola 1º ciclo
    32    Comercial 2º ciclo
    34    Eclesiástico 1º ciclo
    35    Educação física 1º ciclo
    36    Enfermagem 1º ciclo
    39    Industrial 1º ciclo
    41    Militar 1º ciclo
    42    Normal 1º ciclo
    47    Outros 1º ciclo
    48    Emendativo 1º ciclo
    50    Colegial
    51    Agrícola 2º ciclo
    52    Comercial 2º ciclo
    53    Belas artes 2º ciclo
    54    Eclesiástico 2º ciclo
    55    Educação física 2º ciclo
    56    Enfermagem 2º ciclo
    58    Estatística 2º ciclo
    59    Industrial 2º ciclo
    61    Militar 2º ciclo
    62    Normal 2º ciclo
    65    Serviço social 2º ciclo
    67    Outros 2º ciclo
    70    Administração
    71    Agronomia
    72    Arquitetura
    73    Belas artes superior
    74    Ciências sociais
    75    Filosofia
    76    Geografia ou história
    77    História natural
    78    Letras
    79    Matemática, física e química
    80    Pedagogia
    81    Contabilidade ou atuária
    82    Economia
    83    Direito
    84    Eclesiástico superior
    85    Educação física superior
    86    Enfermagem superior
    87    Engenharia
    88    Estatística superior
    89    Farmácia ou bioquímica
    90    Medicina
    91    Militar superior
    92    Odontologia
    93    Psicologia
    94    Química industrial
    95    Serviço social superior
    96    Veterinária
    97    Outros superiores
    98    Grau indeterminado
    99    Nenhum
    ---------------------------------
    '''

    #Coluna de anos de escolaridade
    df[C_ANOS_ESC] = pd.NA

    #Anos base de estudo para cada grau
    yearsStage = {
        '1': 0,
        '2': MAX_EF_AI,
        '3': MAX_EF_AF,
        '4': MAX_EM
    }
    #Anos de estuda para as séries
    yearsSeries = {
        '1': 1,
        '2': 1,
        '3': 2,
        '4': 3,
        '5': 4,
        '6': 5,
        '7': 0,
        '8': 0,
        '9': 0,
        '0': 0,
    }

    df[C_ANOS_ESC] = df.V038.map(yearsStage) + df.V037.map(yearsSeries)
    df.loc[df.V038 == '5', C_ANOS_ESC] = 0
    #Aplicar teto para os graus de escolaridade
    #Elementar: 4 anos; Médio 1º Ciclo: 8; Médio 2º Ciclo: 11; Superior: 15
    df.loc[(df[C_ANOS_ESC] > MAX_EF_AI) & (df.V038 == '1'), C_ANOS_ESC] = MAX_EF_AI
    df.loc[(df[C_ANOS_ESC] > MAX_EF_AF) & (df.V038 == '2'), C_ANOS_ESC] = MAX_EF_AF
    df.loc[(df[C_ANOS_ESC] > MAX_EM)    & (df.V038 == '3'), C_ANOS_ESC] = MAX_EM
    df.loc[(df[C_ANOS_ESC] > MAX_ES)    & (df.V038 == '4'), C_ANOS_ESC] = MAX_ES
    df[C_ANOS_ESC] = df[C_ANOS_ESC].astype('UInt8')

    #Coluna de conclusão de etapas
    #Preparar algumas variáveis
    V039i = df.V039.astype('UInt16')
    df[C_ETAPA_CONC] = pd.NA
    
    df.loc[(V039i == 99) & (df.V035 == '2'), C_ETAPA_CONC] = NAO_CONCLUIU_ANALF
    df.loc[(V039i == 99) & (df.V035 == '1'), C_ETAPA_CONC] = NAO_CONCLUIU_ALFA
    df.loc[V039i.between(10, 28), C_ETAPA_CONC] = EF_AI
    df.loc[V039i.between(30, 49), C_ETAPA_CONC] = EF_AF
    df.loc[V039i.between(50, 69), C_ETAPA_CONC] = EM
    df.loc[V039i.between(70, 97), C_ETAPA_CONC] = ES

    #TODO tratar os casos de não declarados
    #df.V039 == 0

    df[C_ETAPA_CONC] = df[C_ETAPA_CONC].astype(CAT_E_TYPES)

def educacao_1980(df):
    '''
    V519 - Sabe ler e escrever
    ==========================
    2    sim
    4    esqueceu
    6    não sabe
    9    sem declaração
         não aplicável
    --------------------------
    
    V520 - Série que frequenta
    ==========================
    0    nenhuma
    1    1ª série
    2    2ª série
    3    3ª série
    4    4ª série
    5    5ª série
    6    6ª série
    7    7ª série
    8    8ª série
    9    sem declaração
         não aplicável
    --------------------------

    V521 - Grau que frequenta
    =========================
    0    nenhuma
    1    primário
    2    ginasial médio
    3    1º grau
    4    2º grau
    5    colegial médio
    6    supletivo 1º grau
    7    supletivo 2º grau
    8    superior
    9    sem declaração
         não aplicável
    -------------------------

    V522 - Curso não seriado frequentado
    ========================
    0    nenhum
    1    pré-escolar
    2    alfabetização para adultos
    3    supletivo 1º grau
    4    supletivo 2º grau
    5    supletivo 1º grau de TV
    6    supletivo 2º grau de TV
    7    vestibular
    8    mestrado ou doutorado
    9    sem declaração
         não aplicável
    ------------------------

    V523 - Última série concluída
    =============================
    0    nenhuma
    1    1ª série
    2    2ª série
    3    3ª série
    4    4ª série
    5    5ª série
    6    6ª série
    7    7ª série
    8    8ª série
    9    sem declaração
         não aplicável
    -----------------------------
    
    V524 - Grau da última série concluída
    =====================================
    0    nenhuma
    1    alfabetização de adultos
    2    primário ou elementar
    3    ginasial ou médio 1o. Ciclo
    4    1o. Grau
    5    2o. Grau
    6    colegial ou médio 2o. Ciclo
    7    superior
    8    mestrado ou doutorado
    9    sem declaração
         não aplicável
    -------------------------------------

    V525 - Tipo do curso mais elevado concluído
    ===========================================
    00    sem curso concluído
    01    primário
    02    educação especial
    03    agrícola elementar
    04    comercial elementar
    05    artesanal elementar
    06    saúde elementar
    07    militar elementar
    08    outros elementares
    10    1º grau
    11    especial - 1º grau
    12    agrícola - 1º ciclo
    13    administração - 1º ciclo
    14    contabilidade -  1º ciclo
    15    básico comercial
    16    eletricidade - 1º grau
    17    mecânica - 1º ciclo
    18    artesanato básico
    19    enfermagem - 1º ciclo
    20    farmácia - 1º grau
    21    militar - 1º ciclo
    22    normal - 1º ciclo
    23    outros - 1º grau
    24    colegial
    25    educação especial - 2º grau
    26    agrícola - 2º ciclo
    27    administração - 2º grau
    28    contabilidade - 2º ciclo
    29    estatística - 2º grau
    30    secretariado - 2º grau
    31    comercial - 2º ciclo
    32    desenho - 2º grau
    33    eletricidade - 2º grau
    34    mecânica - 2º grau
    35    química - 2º grau
    36    industrial - 2º ciclo
    37    enfermagem - 2º ciclo
    38    patologia - 2º ciclo
    39    nutrição - 2º grau
    40    militar - 2º ciclo
    41    normal - 2º grau
    42    outros - 2º ciclo
    43    biologia - superior
    44    educação física - superior
    45    enfermagem - superior
    46    farmácia - superior
    47    medicina - superior
    48    dentista
    49    nutrição - superior
    50    arquitetura - superior
    51    ciências - bacharelado
    52    computação - bacharelado
    53    engenharia - superior
    54    eletricidade - superior
    55    engenharia - superior
    56    química - superior
    57    topografia - superior
    58    estatística - superior
    59    física - bacharelado
    60    geologia - superior
    61    matemática - bacharelado
    62    química - bacharelado
    63    astronomia - superior
    64    agrícola - superior
    65    veterinária - superior
    66    agrimensura - superior
    67    administração - superior
    68    biblioteconomia - superior
    69    contabilidade - superior
    70    economia - superior
    71    antropologia - superior
    72    comunicação - superior
    73    direito - superior
    74    filosofia - bacharelado
    75    geografia - bacharelado
    76    história - bacharelado
    77    educação - superior
    78    psicologia - bacharelado
    79    serviço social - superior
    80    sacerdote - superior
    81    museologia - superior
    82    lingüística - superior
    83    belas artes - superior
    84    academia militar
    85    diplomacia - superior
    86    medicina - mestrado
    87    biologia - mestrado
    88    engenharia - mestrado
    89    computação - mestrado
    90    agronomia - mestrado
    91    administração - mestrado
    92    economia - mestrado
    93    direito - mestrado
    94    educação - mestrado
    95    biblioteconomia - mestrado
    96    artes - mestrado
    99    ignorado
    -------------------------------------------
    '''

    #Coluna de anos de escolaridade
    df[C_ANOS_ESC] = pd.NA

    #Anos base de estudo para cada grau concluído
    yearsStageV524 = {
        '1': 0,
        '2': 0,
        '3': MAX_EF_AI,
        '4': 0,
        '5': MAX_EF_AF,
        '6': MAX_EF_AF,
        '7': MAX_EM,
        '8': MAX_ES,
    }
    #Anos base de estudo para cada grau frequentado
    yearsStageV521 = {
        '0': 0,
        '1': 0,
        '2': MAX_EF_AI,
        '3': 0,
        '4': MAX_EF_AF,
        '5': MAX_EF_AF,
        '6': 0,
        '7': MAX_EF_AF,
        '8': MAX_EM,
    }
    #Anos de estuda para as séries
    yearsSeries = {
        '0': 0,
        '1': 1,
        '2': 2,
        '3': 3,
        '4': 4,
        '5': 5,
        '6': 6,
        '7': 7,
        '8': 8,
        '9': 0,
    }

    df[C_ANOS_ESC] = df.V524.map(yearsStageV524) + df.V523.map(yearsSeries)
    df.loc[df.V524 == '0', C_ANOS_ESC] = 0

    # Cursos não seriados
    #df.loc[df.V522.isin({'1', '2', '3', '5'}), C_ANOS_ESC] = 0
    df.loc[df.V522.isin({'4', '6'}) & (df[C_ANOS_ESC] < MAX_EF_AF), C_ANOS_ESC] = MAX_EF_AF
    df.loc[(df.V522 == '7') & (df[C_ANOS_ESC] < MAX_EM), C_ANOS_ESC] = MAX_EM
    df.loc[(df.V522 == '8') & (df[C_ANOS_ESC] < MAX_ES), C_ANOS_ESC] = MAX_ES

    #Nota, p. 22 do Manual do Recenseador: "Para a maioria das pessoas que
    #frequentam escola, estará assinalado, nas variáveis V523 e V524, o
    #valor 0 -- Nenhum, salvo para as pessoas que tenham concluído ou
    #interrompido algum curso, mas estejam fazendo um outro curso de mesmo grau ou grau
    # ou grau inferior ao curso concluído

    #NOTE o código abaixo é temporário, excluir após refazer o parquet
    df.V521 = df.V521.astype('string')

    # df.loc[df.V524 == '0', C_ANOS_ESC] = (df.V521.map(yearsStageV521) + df.V520.map(yearsSeries) - 1).apply(lambda s: max(s, 0))
    anos_esc_tmp = (df.V521.map(yearsStageV521) + df.V520.map(yearsSeries) - 1).apply(lambda s: max(s, 0))
    df.loc[:, C_ANOS_ESC] = pd.concat([anos_esc_tmp, df[C_ANOS_ESC]], axis=1).max(axis=1)

    #Aplicar teto para os graus de escolaridade
    df.loc[(df[C_ANOS_ESC] > MAX_EF_AI) & ((df.V521 == '1') | df.V524.isin({'1', '2'})), C_ANOS_ESC] = MAX_EF_AI
    df.loc[(df[C_ANOS_ESC] > MAX_EF_AF) & (df.V521.isin({'2', '3', '6'}) | df.V524.isin({'3', '4'})), C_ANOS_ESC] = MAX_EF_AF
    df.loc[(df[C_ANOS_ESC] > MAX_EM) & (df.V521.isin({'4', '5', '7'}) | df.V524.isin({'5', '6'})), C_ANOS_ESC] = MAX_EM
    df.loc[(df[C_ANOS_ESC] > MAX_ES) & ((df.V521 == '8') | df.V524.isin({'7', '8'})), C_ANOS_ESC] = MAX_ES

    V525i = df.V525.astype('UInt16')
    df.loc[V525i.between(1,   8) & (df[C_ANOS_ESC] < MAX_EF_AI), C_ANOS_ESC] = MAX_EF_AI
    df.loc[V525i.between(10, 23) & (df[C_ANOS_ESC] < MAX_EF_AF), C_ANOS_ESC] = MAX_EF_AF
    df.loc[V525i.between(24, 42) & (df[C_ANOS_ESC] < MAX_EM),    C_ANOS_ESC] = MAX_EM
    df.loc[V525i.between(43, 96) & (df[C_ANOS_ESC] < MAX_ES),    C_ANOS_ESC] = MAX_ES

    df[C_ANOS_ESC] = df[C_ANOS_ESC].astype('UInt8')

    #Coluna de conclusão de etapas
    conc_tmp0 = pd.Series(data=pd.NA, index=df.index)
    conc_tmp0[(V525i == 0) & df.V519.isin({'4', '6'})] = NAO_CONCLUIU_ANALF
    conc_tmp0[(V525i == 0) & (df.V519 == '2')] = NAO_CONCLUIU_ALFA
    conc_tmp0[(V525i == 0) & (df.V519 == '9')] = NAO_CONCLUIU_SEM_DECLARACAO
    conc_tmp0[V525i.between(1,   8)] = EF_AI
    conc_tmp0[V525i.between(10, 23)] = EF_AF
    conc_tmp0[V525i.between(24, 42)] = EM
    conc_tmp0[V525i.between(43, 96)] = ES

    conc_tmp1 = pd.Series(data=pd.NA, index=df.index)
    #TODO incluir supletivo 1º grau
    conc_tmp1[(df.V521 == '2') 
              | ((df.V521 == '3') & df.V520.isin({'5', '6', '7', '8'})) 
              | (df.V524 == '3') 
              | ((df.V524 == '4') & df.V523.isin({'4', '5', '6', '7'}))] = EF_AI

    conc_tmp1[df.V521.isin({'4', '5', '7'})
              | df.V521.isin({'4', '5'})
              | ((df.V524 == '3') & (df.V523 == '4')) 
              | ((df.V524 == '4') & (df.V523 == '8')) 
              | (df.V524.isin({'5', '6'}))] = EF_AF

    conc_tmp1[(df.V521 == '8')
              | (df.V522 == '7')
              | (df.V524.isin({'5', '6'}) & df.V523.isin({'3', '4'})) 
              | (df.V524 == '7')] = EM

    conc_tmp1[(df.V522 == '8')
              | ((df.V524 == '7') & df.V523.isin({'5', '6', '7', '8'})) 
              | (df.V524 == '8')] = ES
               
    df[C_ETAPA_CONC] = pd.concat([conc_tmp0, conc_tmp1], axis=1).astype(CAT_E_TYPES).max(axis=1)

    #Coluna etapa frequentada
    #Os supletivos seriados foram classificados no ensino regular
    #exceto aqueles casos sem declaração da série que frequentada
    #nestes casos foram considerados `S1` ou `S2`
    #Os sem declaração da série frequentada foram considerados
    #numa categoria à parte
    freq = pd.Series(data=pd.NA, index=df.index)
    freq[(df.V521 == '0') & (df.V522 == '0')] = NAO_FREQUENTA
    freq[df.V521.isin({'1', '2', '3'}) & (df.V520 == '9')] = EF_SD
    freq[df.V521.isin({'1', '3', '6'}) & (df.V520 == '1')] = EF_1
    freq[df.V521.isin({'1', '3', '6'}) & (df.V520 == '2')] = EF_2
    freq[df.V521.isin({'1', '3', '6'}) & (df.V520 == '3')] = EF_3

    freq[((df.V521 == '1') & df.V520.isin({'4', '5', '6', '7', '8'}))
       | (df.V521.isin({'3', '6'}) & (df.V520 == '4'))] = EF_4

    freq[(df.V521.isin({'3', '6'}) & (df.V520 == '5'))
       | ((df.V521 == '2') & (df.V520 == '1'))] = EF_5

    freq[(df.V521.isin({'3', '6'}) & (df.V520 == '6'))
       | ((df.V521 == '2') & (df.V520 == '2'))] = EF_6

    freq[(df.V521.isin({'3', '6'}) & (df.V520 == '7'))
       | ((df.V521 == '2') & (df.V520 == '3'))] = EF_7

    freq[(df.V521.isin({'3', '6'}) & (df.V520 == '8'))
       | ((df.V521 == '2') & (df.V520 == '4'))] = EF_8

    freq[df.V521.isin({'4', '5'}) & (df.V520 == '9')] = EM_SD
    freq[df.V521.isin({'4', '5', '7'}) & (df.V520 == '1')] = EM_1
    freq[df.V521.isin({'4', '5', '7'}) & (df.V520 == '2')] = EM_2
    freq[df.V521.isin({'4', '5', '7'}) & df.V520.isin({'3', '4', '5', '6', '7', '8'})] = EM_3

    freq[(df.V521 == '8') & (df.V520 == '9')] = ES_SD
    freq[(df.V521 == '8') & (df.V520 == '1')] = ES_1
    freq[(df.V521 == '8') & (df.V520 == '2')] = ES_2
    freq[(df.V521 == '8') & (df.V520 == '3')] = ES_3
    freq[(df.V521 == '8') & (df.V520 == '4')] = ES_4
    freq[(df.V521 == '8') & (df.V520 == '5')] = ES_5
    freq[(df.V521 == '8') & df.V520.isin({'6', '7', '8'})] = ES_6

    freq[df.V522 == '1'] = PE
    freq[df.V522 == '2'] = AA
    freq[df.V522.isin({'3', '5'}) | ((df.V521 == '6') & (df.V520 == '9'))] = S1
    freq[df.V522.isin({'4', '6'}) | ((df.V521 == '7') & (df.V520 == '9'))] = S2
    freq[df.V522 == '7'] = V
    freq[df.V522 == '8'] = PG

    df[C_FREQ] = freq.astype(CAT_F_TYPES)

def educacao_1991(df):
    '''
    V0323 - Alfabetização
    =====================
    1   Sabe ler e escrever
    2   Não sabe
        Pessoas com menos de 5 anos
    ---------------------

    V0324 - Série que frequenta
    ===========================
    0   Nenhuma
    1   1ª Série
    2   2ª Série
    3   3ª Série
    4   4ª Série
    5   5ª Série
    6   6ª Série
    7   7ª Série
    8   8ª Série
        Pessoas com menos de 5 anos
    ---------------------------

    V0325 - Grau que frequenta em curso seriado
    ===========================================
    0   Nenhum
    1   1º grau
    2   2º grau
    3   Superior
    4   Supletivo - 1º grau
    5   Supletivo - 2º grau
        Pesoas com menos de 5 anos
    -------------------------------------------

    V0326 - Grau que frequenta em curso não seriado
    ===============================================
    0   Nenhum
    1   Pré-escolar
    2   Curso de alfabetização de adultos
    3   Supletivo não seriado - 1º grau
    4   Supletivo não seriado - 2º grau
    5   Pré-vestibular
    6   Mestrado ou Doutorado
        Pessoas com menos de 5 anos
    -----------------------------------------------

    V0327 - Última série concluída com aprovação
    ============================================
    0   Nenhuma
    1   1ª Série
    2   2ª Série
    3   3ª Série
    4   4ª Série
    5   5ª Série
    6   6ª Série
    7   7ª Série
    8   8ª Série
    9   Nunca frequentou
        Pessoas com menos de 5 anos
    --------------------------------------------

    V0328 - Grau da última série concluída com aprovação
    ====================================================
    0   Nenhum
    1   Curso de alfabetização de adultos
    2   Primário ou elementar
    3   Ginasial ou médio 1º ciclo
    4   1º grau
    5   2º grau
    6   Colegial ou médio 2º ciclo
    7   Superior
    8   Mestrado ou Doutorado
        Pessoas com menos de 5 anos
    ----------------------------------------------------

    V3241 - Anos de estudo
    ======================
    00      Sem instrução
    30      Alfabetização de adultos
    01 a 16 Número de anos de estudo
    17      17 anos ou mais de estudos
    20      Não determinado
            Pessoas com menos de 5 anos
    ----------------------

    V0329 - Curso concluído
    =======================
    00   Nenhum curso
    01   Primário ou Elementar (Ensino Geral)
    02   Primário ou Elementar (Educação Especial)
    03   Primário ou Elementar (Agrícola)
    04   Primário ou Elementar (Administração)
    05   Primário ou Elementar (Industrial)
    06   Primário ou Elementar (Saúde)
    07   Primário ou Elementar (Militar)
    08   Primário ou Elementar (Outros)
    10   Ensino Geral - 1º grau
    11   Educação Especial - 1º grau
    12   Agrícola - 1º grau
    13   Administração - 1º grau
    14   Contabilidade - 1º grau
    15   Outro - 1º grau - Comercial
    16   Eletrotécnica ou Eletrônica - 1º grau
    17   Mecânica - 1º grau
    18   Outro -1º grau - Industrial
    19   Enfermagem - 1º grau
    20   Outros - 1º grau - Saúde
    21   Militar - 1º grau
    22   Normal - 1º grau
    23   Outros - 1º grau
    24   Ensino Geral - 2º grau
    25   Educação Especial - 2º grau
    26   Agrícola - 2º grau
    27   Administração - 2º grau
    28   Contabilidade - 2º grau
    29   Estatística - 2º grau
    30   Secretariado - 2º grau
    31   Outros - 2º grau - Comercial
    32   Desenho - 2º grau
    33   Eletrotécnica ou Eletrônica - 2º grau
    34   Mecânica - 2º grau
    35   Química - 2º grau
    36   Outros - 2º grau - Industrial
    37   Enfermagem - 2º grau
    38   Laboratorista de análise clínica - 2º grau
    39   Outros - 2º grau - Saúde
    40   Militar - 2º grau
    41   Normal - 2º grau
    42   Outros - 2º grau
    43   Biologia (inclui: Biomedicina, Biologia, Ciências Biológicas e afins)
    44   Educação Física (inclui: Técnico de desportos)
    45   Enfermagem
    46   Farmácia (inclui: Farmácia Bioquímica e afins)
    47   Medicina
    48   Odontologia
    49   Outros da Biologia (inclui: Nutrição, Fisioterapia, fonoaudiologia, Terapia Ocupacional)
    50   Arquitetura e Urbanismo
    51   Ciências Exatas
    52   Ciências da Computação (inclui: Processamento de Dados, Engenharia de Sistemas, Informática e afins)
    53   Engenharia Civil (inclui: Engenharia de Edificações, Pontes, Estradas, etc.)
    54   Engenharia Elétrica e Eletrônica (inclui: Eletrotécnica, Telecomunicações, Comunicações, etc.)
    55   Engenharia Mecânica (inclui: Engenharia Naval, Aeronáutica, Metalúrgica, Mecânica de Automóveis, etc.)
    56   Engenharia Química e Química Industrial
    57   Engenharia não classificada ou mal definida
    58   Estatística
    59   Física
    60   Geologia
    61   Matemática
    62   Química
    63   Outros da Tecnologia (exclusive Engenharia)
    64   Agronomia (inclui: Ciências Agrícolas, Engenharia Agrícola e Engenharia Florestal)
    65   Medicina Veterinária
    66   Outros (inclui: Agrimensura, Engenharia de Pesca, Zootecnia, Fitotecnia) - Agrárias
    67   Administração (inclui: Secretariado, Administração de Empresas, Administração Pública)
    68   Biblioteconomia (inclui: Arquivologia)
    69   Ciências Contábeis e Atuariais
    70   Ciências Econômicas
    71   Ciências e Estudos Sociais (inclui: Sociologia, Ciências Políticas e Sociais, Antropologia, etc.)
    72   Comunicação Social (inclui: Editoração, Jornalismo, Rádio e TV, Publicidade, Turismo)
    73   Direito (inclui: Relações Internacionais)
    74   Filosofia
    75   Geografia
    76   História
    77   Pedagogia (inclui: Administração Escolar, Educação Especial, Orientação Educacional, Formação de Professores)
    78   Psicologia
    79   Serviço Social
    80   Teologia
    81   Outros de Humanas (inclui: Museologia, Arqueologia)
    82   Letras (inclui: Tradutor e Intérprete)
    83   Artes (inclui: Cinema, Comunicação Visual, Dança, História da Arte, Desenho, Música, Teatro, etc.)
    84   Defesa Nacional (Militar)
    85   Outros Cursos de Grau Superior
    86   Mestrado ou Doutorado - Medicina
    87   Mestrado ou Doutorado - Outros (Biologia)
    88   Mestrado ou Doutorado - Engenharia
    89   Mestrado ou Doutorado - Outros (Ciências Tecnológicas)
    90   Mestrado ou Doutorado - Ciências Agrárias
    91   Mestrado ou Doutorado - Administração
    92   Mestrado ou Doutorado - Ciências Econômicas, Contábeis, etc.
    93   Mestrado ou Doutorado - Direito
    94   Mestrado ou Doutorado - Pedagogia
    95   Mestrado ou Doutorado - Outros (Ciências Humanas e Sociais)
    96   Mestrado ou Doutorado - Letras e Artes
    97   Mestrado ou Doutorado - (área não especificada)
    -----------------------
    '''

    #Coluna de anos de escolaridade
    df[C_ANOS_ESC] = pd.NA
    
    V3241 = df.V3241
    df.loc[V3241.between(0, 17), C_ANOS_ESC] = V3241[V3241.between(0, 17)]
    df.loc[V3241 == 30, C_ANOS_ESC] = 0
    df[C_ANOS_ESC] = df[C_ANOS_ESC].astype('UInt8')

    #Coluna etapa frequentada
    freq = pd.Series(data=pd.NA, index=df.index)

    freq[(df.V0325 == '0') & (df.V0326 == '0')] = NAO_FREQUENTA
    f_EF = df.V0325.isin({'1', '4'})
    freq[f_EF & (df.V0324 == 0)] = EF_SD
    freq[f_EF & (df.V0324 == 1)] = EF_1
    freq[f_EF & (df.V0324 == 2)] = EF_2
    freq[f_EF & (df.V0324 == 3)] = EF_3
    freq[f_EF & (df.V0324 == 4)] = EF_4
    freq[f_EF & (df.V0324 == 5)] = EF_5
    freq[f_EF & (df.V0324 == 6)] = EF_6
    freq[f_EF & (df.V0324 == 7)] = EF_7
    freq[f_EF & (df.V0324 == 8)] = EF_8

    f_EM = df.V0325.isin({'2', '5'})
    freq[f_EM & (df.V0324 == 0)] = EM_SD
    freq[f_EM & (df.V0324 == 1)] = EM_1
    freq[f_EM & (df.V0324 == 2)] = EM_2
    freq[f_EM & (df.V0324 >= 3)] = EM_3

    f_ES = df.V0325 == '3'
    freq[f_ES & (df.V0324 == 0)] = ES_SD
    freq[f_ES & (df.V0324 == 1)] = ES_1
    freq[f_ES & (df.V0324 == 2)] = ES_2
    freq[f_ES & (df.V0324 == 3)] = ES_3
    freq[f_ES & (df.V0324 == 4)] = ES_4
    freq[f_ES & (df.V0324 == 5)] = ES_5
    freq[f_ES & (df.V0324 >= 6)] = ES_6

    freq[df.V0326 == '1'] = PE
    freq[df.V0326 == '2'] = AA
    freq[df.V0326 == '3'] = S1
    freq[df.V0326 == '4'] = S2
    freq[df.V0326 == '5'] = V
    freq[df.V0326 == '6'] = PG

    df[C_FREQ] = freq.astype(CAT_F_TYPES)

def educacao_2000(df):
    '''
    V0428 - Sabe ler e escrever
    ---------------------------
    1   Sabe ler e escrever
    2   Não sabe
    ===========================

    V0429 - Frequenta escola ou creche
    ----------------------------------
    1   Sim, rede particular
    2   Sim, rede pública
    3   Não, já frequentou
    4   Nunca frequentou
    ==================================

    V0430 - Curso que frequenta
    ---------------------------
    01  Creche
    02  Pré-escolar
    03  Classe de alfabetização
    04  Alfabetização de adultos
    05  Ensino fundamental ou 1º grau – regular seriado
    06  Ensino fundamental ou 1º grau – regular não-seriado
    07  Supletivo (ensino fundamental ou 1º grau)
    08  Ensino médio ou 2º grau – regular seriado
    09  Ensino médio ou 2º grau – regular não-seriado
    10  Supletivo (ensino médio ou 2º grau)
    11  Pré-vestibular
    12  Superior – graduação
    13  Mestrado ou doutorado
        Para os não estudantes
    ===========================

    V0431 - Série que frequenta
    ---------------------------
    1   Primeira Série
    2   Segunda Série
    3   Terceira Série
    4   Quarta Série
    5   Quinta Série
    6   Sexta Série
    7   Sétima Série
    8   Oitava Série
    9   Curso não seriado
        Para os não estudantes
    ===========================

    V0432 - Curso mais elevado que frequentou, concluindo pelo menos uma série
    --------------------------------------------------------------------------
    1   Alfabetização de adultos
    2   Antigo primário
    3   Antigo ginásio
    4   Antigo clássico, científico, etc.
    5   Ensino fundamental ou 1º grau
    6   Ensino médio ou 2º Grau
    7   Superior – graduação
    8   Mestrado ou doutorado
    9   Nenhum
        Para os estudantes
    ==========================================================================

    V0433 - Última série concluída com aprovação
    --------------------------------------------
    01  Primeira Série
    02  Segunda Série
    03  Terceira Série
    04  Quarta Série
    05  Quinta Série
    06  Sexta Série
    07  Sétima Série
    08  Oitava Série
    09  Curso não seriado
    10  Nenhuma
        Para os estudantes
    ============================================

    V0434 - Concluiu o curso no qual estudou
    ----------------------------------------
    1   Sim
    2   Não
        Para os estudantes
    ========================================

    V4355 - Código do curso mais elevado concluído
    ----------------------------------------------
    01  Outros cursos de Graduação
    02  Não superior
    09  Outros cursos de Mestrado ou Doutorado
    11  Agronomia - Graduação
    12  Medicina Veterinária - Graduação
    13  Outros de Ciências Agrárias - Graduação
    19  Ciências Agrárias - Mestrado ou Doutorado
    21  Biologia - Graduação
    22  Educação Física - Graduação
    23  Enfermagem - Graduação
    24  Farmácia - Graduação
    25  Medicina - Graduação
    26  Odontologia - Graduação
    27  Outros de Ciências Biológicas e da Saúde - Graduação
    28  Medicina - Mestrado ou Doutorado
    29  Outros de Ciências Biológicas e da Saúde - Mestrado ou Doutorado
    31  Arquitetura e Urbanismo - Graduação
    32  Ciências - Graduação
    33  Ciências da Computação - Graduação
    34  Engenharia Civil - Graduação
    35  Engenharia Elétrica e Eletrônica - Graduação
    36  Engenharia Mecânica - Graduação
    37  Engenharia Química e Industrial - Graduação
    38  Outros cursos de Engenharia - Graduação
    39  Engenharia - Mestrado ou Doutorado
    41  Estatística - Graduação
    42  Física - Graduação
    43  Geologia - Graduação
    44  Matemática - Graduação
    45  Química - Graduação
    46  Outros de Ciências Exatas e Tecnológicas, exclusive Engenharia - Graduação
    49  Outros de Ciências Exatas e Tecnológicas, exclusive Engenharia - Mestrado ou doutorado
    51  Administração - Graduação
    52  Biblioteconomia - Graduação
    53  Ciências Contábeis e Atuariais - Graduação
    54  Ciências Econômicas - Graduação
    55  Ciências e Estudos Sociais - Graduação
    56  Comunicação Social - Graduação
    57  Direito - Graduação
    58  Filosofia - Graduação
    59  Formação Professores Disciplinas Especiais - Graduação
    61  Geografia - Graduação
    62  História - Graduação
    63  Pedagogia - Graduação
    64  Propaganda e Marketing - Graduação
    65  Psicologia - Graduação
    66  Serviço Social - Graduação
    67  Teologia - Graduação
    68  Outros de Ciências Humanas e Sociais - Graduação
    75  Administração - Mestrado ou Doutorado
    76  Ciências Econômicas e Contábeis - Mestrado ou Doutorado
    77  Direito - Mestrado ou Doutorado
    78  Pedagogia - Mestrado ou Doutorado
    79  Outros de Ciências Humanas e Sociais - Mestrado ou Doutorado
    81  Letras - Graduação
    82  Artes - Graduação
    83  Outros de Letras e Artes - Graduação
    89  Letras e Artes - Mestrado ou Doutorado
    91  Militar
    ==============================================

    V4300 - Anos de estudo
    ----------------------
    00  Sem instrução ou menos de 1 ano
    01  1 ano
    02  2 anos
    03  3 anos
    04  4 anos
    05  5 anos
    06  6 anos
    07  7 anos
    08  8 anos
    09  9 anos
    10  10 anos
    11  11 anos
    12  12 anos
    13  13 anos
    14  14 anos
    15  15 anos
    16  16 anos
    17  17 anos ou mais
    20  Não determinado
    30  Alfabetização de adultos
    ======================
    '''

    #Coluna de anos de escolaridade
    df[C_ANOS_ESC] = pd.NA
    
    V4300i = df.V4300.astype('UInt8')
    df.loc[V4300i.between(0, 17), C_ANOS_ESC] = V4300i[V4300i.between(0, 17)]
    df.loc[V4300i == 30, C_ANOS_ESC] = 0
    df.loc[(V4300i == 8) & df.V0432.isin({'4', '6'}) & (df.V0434 == '1'), C_ANOS_ESC] = 11 # Possível erro na variável V4300 - anos de estudo
    df[C_ANOS_ESC] = df[C_ANOS_ESC].astype('UInt8')

    #Coluna etapa frequentada
    freq = pd.Series(data=pd.NA, index=df.index)

    V0431i = df.V0431.astype('UInt8')

    freq[df.V0429.isin({'3', '4'})] = NAO_FREQUENTA
    f_EF = df.V0430.isin({'05', '06', '07'})
    freq[f_EF & (V0431i == 1)] = EF_1
    freq[f_EF & (V0431i == 2)] = EF_2
    freq[f_EF & (V0431i == 3)] = EF_3
    freq[f_EF & (V0431i == 4)] = EF_4
    freq[f_EF & (V0431i == 5)] = EF_5
    freq[f_EF & (V0431i == 6)] = EF_6
    freq[f_EF & (V0431i == 7)] = EF_7
    freq[f_EF & (V0431i == 8)] = EF_8
    freq[f_EF & (V0431i == 9)] = EF_SD

    f_EM = df.V0430.isin({'08', '09', '10'})
    freq[f_EM & (V0431i == 1)] = EM_1
    freq[f_EM & (V0431i == 2)] = EM_2
    freq[f_EM & V0431i.between(3, 8)] = EM_3
    freq[f_EM & (V0431i == 9)] = EM_SD

    f_ES = df.V0430 == '12'
    freq[f_ES & (V0431i == 1)] = ES_1
    freq[f_ES & (V0431i == 2)] = ES_2
    freq[f_ES & (V0431i == 3)] = ES_3
    freq[f_ES & (V0431i == 4)] = ES_4
    freq[f_ES & (V0431i == 5)] = ES_5
    freq[f_ES & (V0431i >= 6)] = ES_6
    freq[f_ES & (V0431i == 9)] = ES_SD

    freq[df.V0430.isin({'01', '02', '03'})] = PE
    freq[df.V0430 == '04'] = AA
    freq[(df.V0430 == '07') & (V0431i == 9)] = S1
    freq[(df.V0430 == '10') & (V0431i == 9)] = S2
    freq[df.V0430 == '11'] = V
    freq[df.V0430 == '13'] = PG

    df[C_FREQ] = freq.astype(CAT_F_TYPES)

    #Coluna de conclusão de etapas
    conc_tmp0 = pd.Series(data=pd.NA, index=df.index)
    conc_tmp0[(df.V0432 == '9') & (df.V0428 == '2')] = NAO_CONCLUIU_ANALF
    conc_tmp0[(df.V0432 == '9') & (df.V0428 == '1')] = NAO_CONCLUIU_ALFA
    conc_tmp0[((df.V0432 == '2') & (df.V0434 == '1'))
            | ((df.V0432 == '3') & (df.V0434 == '2'))
            | ((df.V0432 == '5') & df.V0433.isin({'04', '05', '06', '07'}))] = EF_AI
    conc_tmp0[(df.V0432.isin({'3', '5'}) & (df.V0434 == '1'))
            | (df.V0432.isin({'4', '6'}) & (df.V0434 == '2'))] = EF_AF
    conc_tmp0[(df.V0432.isin({'4', '6'}) & (df.V0434 == '1'))
           | ((df.V0432 == '7') & (df.V0434 == '2'))] = EM
    conc_tmp0[((df.V0432 == '7') & (df.V0434 == '1'))
             | (df.V0432 == '8')] = ES

    conc_tmp1 = pd.Series(data=pd.NA, index=df.index)
    conc_tmp1[df[C_FREQ].isin({EF_1, EF_2, EF_3, EF_4, PE, AA}) & (df.V0428 == '1')] = NAO_CONCLUIU_ANALF
    conc_tmp1[df[C_FREQ].isin({EF_1, EF_2, EF_3, EF_4, PE, AA}) & (df.V0428 == '2')] = NAO_CONCLUIU_ALFA
    conc_tmp1[df[C_FREQ].isin({EF_5, EF_6, EF_7, EF_8, S1})]= EF_AI
    conc_tmp1[df[C_FREQ].isin({EM_1, EM_2, EM_3, S2})]= EF_AF
    conc_tmp1[df[C_FREQ].isin({ES_1, ES_2, ES_3, ES_4, ES_5, ES_6, V})]= EM
    conc_tmp1[df[C_FREQ] == PG] = ES

    df[C_ETAPA_CONC] = pd.concat([conc_tmp0, conc_tmp1], axis=1).astype(CAT_E_TYPES).max(axis=1)

def educacao_2010(df):
    '''
    V0627 - Sabe ler e escrever
    ---------------------------
    1   Sim
    2   Não
        Pessoas menores de 5 anos de idade
    ===========================

    V0628 - Frequenta escola ou creche
    ----------------------------------
    1   Sim, pública
    2   Sim, particular
    3   Não, já frequentou
    4   Não, nunca frequentou
    ==================================

    V0629 - Curso que frequenta
    ---------------------------
    01  Creche
    02  Pré-escolar (maternal e jardim da infância)
    03  Classe de alfabetização - CA
    04  Alfabetização de jovens e adultos
    05  Regular do ensino fundamental
    06  Educação de jovens e adultos - EJA ou supletivo do ensino fundamental
    07  Regular do ensino médio
    08  Educação de jovens e adultos - EJA ou supletivo do ensino médio
    09  Superior de graduação
    10  Especialização de nível superior (mínimo de 360 horas)
    11  Mestrado
    12  Doutorado
        Não frequentavam ou nunca frequentaram escola ou creche
    ===========================

    V0630 - Serie/ano que frequenta 
    --------------------------------
    01  Primeiro ano
    02  Primeira série / Segundo ano
    03  Segunda série / Terceiro ano
    04  Terceira série / Quarto ano
    05  Quarta série / Quinto ano
    06  Quinta série / Sexto ano
    07  Sexta série / Sétimo ano
    08  Sétima série / Oitavo ano
    09  Oitava série / Nono ano
    10  Não seriado
        Não frequentavam, ou que nunca frequentaram escola ou creche, ou quem frequentava o curso Regular do Ensino Médio
    ================================

    V0631 - Série que frequenta
    ---------------------------
    1   Primeira série
    2   Segunda série
    3   Terceira série
    4   Quarta série
    5   Não seriado
        Não frequentavam, ou que nunca frequentaram escola ou creche, ou quem frequentava o curso Regular do Ensino Fundamental
    ===========================
    Nota: apenas para aqueles que frequentavam o Ensino Médio

    V0632 - Conclusão de outro curso superior de graduação
    ------------------------------------------------------
    1   Sim
    2   Não
        Não frequentavam, ou que nunca frequentaram escola ou creche, ou quem frequentava o curso Superior de Graduação
    ======================================================

    V0633 - Curso mais elevado que frequentou
    -----------------------------------------
    01  Creche, pré-escolar (maternal e jardim de infância), classe de alfabetização - CA
    02  Alfabetização de jovens e adultos
    03  Antigo primário (elementar)
    04  Antigo ginásio (médio 1º ciclo)
    05  Ensino fundamental ou 1º grau (da 1ª a 3ª série/ do 1º ao 4º ano)
    06  Ensino fundamental ou 1º grau (4ª série/ 5º ano)
    07  Ensino fundamental ou 1º grau (da 5ª a 8ª série/ 6º ao 9º ano)
    08  Supletivo do ensino fundamental ou do 1º grau
    09  Antigo científico, clássico, etc.....(médio 2º ciclo)
    10  Regular ou supletivo do ensino médio ou do 2º grau
    11  Superior de graduação
    12  Especialização de nível superior ( mínimo de 360 horas )
    13  Mestrado
    14  Doutorado
        Frequentavam escola ou creche ou aqueles que nunca frequentaram.
    =========================================

    V0634 - Conclusão deste curso
    -----------------------------
    1   Sim
    2   Não
        Frequentavam escola ou creche ou nunca frequentaram
    =============================

    V0635 - Espécie do curso mais elevado concluído
    -----------------------------------------------
    1   Superior de graduação
    2   Mestrado
    3   Doutorado
    ===============================================

    V6400 - Nível de instrução
    --------------------------
    1   Sem instrução e fundamental incompleto
    2   Fundamental completo e médio incompleto
    3   Médio completo e superior incompleto
    4   Superior completo
    5   Não determinado
    ==========================
    '''
    #Coluna etapa frequentada
    freq = pd.Series(data=pd.NA, index=df.index)

    V0630i = df.V0630.astype('UInt8')

    freq[df.V0628.isin({'3', '4'})] = NAO_FREQUENTA
    f_EF = df.V0629.isin({'05', '06'})
    freq[f_EF & (V0630i ==  1)] = EF_0
    freq[f_EF & (V0630i ==  2)] = EF_1
    freq[f_EF & (V0630i ==  3)] = EF_2
    freq[f_EF & (V0630i ==  4)] = EF_3
    freq[f_EF & (V0630i ==  5)] = EF_4
    freq[f_EF & (V0630i ==  6)] = EF_5
    freq[f_EF & (V0630i ==  7)] = EF_6
    freq[f_EF & (V0630i ==  8)] = EF_7
    freq[f_EF & (V0630i ==  9)] = EF_8
    freq[f_EF & (V0630i == 10)] = EF_SD

    f_EM = df.V0629 == '07'
    freq[f_EM & (df.V0631 == '1')] = EM_1
    freq[f_EM & (df.V0631 == '2')] = EM_2
    freq[f_EM & df.V0631.isin({'3', '4'})] = EM_3
    freq[f_EM & (df.V0631 == '5')] = EM_SD

    f_ES = df.V0629 == '09'
    #Não há registro da série nesta edição do Censo
    freq[f_ES] = ES_SD

    freq[df.V0629.isin({'01', '02', '03'})] = PE
    freq[df.V0629 == '04'] = AA
    freq[df.V0629 == '06'] = S1
    freq[df.V0629 == '08'] = S2
    #Os outros Censos não são claros quanto a inclusão de especializações de nível superior
    freq[df.V0629.isin({'10', '11', '12'})] = PG

    df[C_FREQ] = freq.astype(CAT_F_TYPES)

    #Coluna de conclusão de etapas
    conc_tmp = pd.Series(data=pd.NA, index=df.index)

    f_nao_concluiu = (
        (df.V6400 == '1')
           & (df[C_FREQ].isin({PE, AA, EF_0, EF_1, EF_2, EF_3, EF_4})
            | (df[C_FREQ].isin({S1, EF_SD}) & (df.idade <= 10))
            | df.V0633.isin({'01', '02', '05'})
            | ((df.V0633 == '03') & (df.V0634 == '2'))
            | (df.V0628 == '4'))
    )

    conc_tmp[f_nao_concluiu] = NAO_CONCLUIU_SEM_DECLARACAO
    conc_tmp[f_nao_concluiu & (df.V0627 == '2')] = NAO_CONCLUIU_ANALF
    conc_tmp[f_nao_concluiu & (df.V0627 == '1')] = NAO_CONCLUIU_ALFA
    #Decisão arbitrária de incluir V0633 == '08' (Supletivo), S1 e
    #V0633 = '06' 4ª Série do EF como EF_AI completo
    conc_tmp[(df.V6400 == '1')
           & (  df[C_FREQ].isin({EF_5, EF_6, EF_7, EF_8})
              | (df[C_FREQ].isin({S1, EF_SD}) & (df.idade > 10))
              | ((df.V0633 == '03') & (df.V0634 == '1'))
              | (df.V0633 == '06')
              | (df.V0633.isin({'04', '07', '08'}) & (df.V0634 == '2')))] = EF_AI
    conc_tmp[df.V6400 == '2'] = EF_AF
    conc_tmp[df.V6400 == '3'] = EM
    conc_tmp[df.V6400 == '4'] = ES

    df[C_ETAPA_CONC] = conc_tmp.astype(CAT_E_TYPES)


AGE = 'idade'
def idade_2010(df):
    df[AGE] = df.V6036
